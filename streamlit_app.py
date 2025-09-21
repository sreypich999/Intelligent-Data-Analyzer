import streamlit as st
import logging
import os
import base64
import re
import sqlite3
from datetime import datetime
import time
import streamlit.components.v1 as components
from fuzzywuzzy import fuzz
import chromadb
from chromadb.config import Settings
from docx import Document
import pandas as pd

from utils.config import DATA_DIR, SUPPORTED_FILE_TYPES, PROFILES_DIR, INSIGHTS_DIR, INTERACTIVE_PLOTS_DIR, REPORT_FORMATS, CHROMA_DB_PATH
from utils.session_manager import SessionManager
from utils.data_loader import DataLoader
from utils.rag_pipeline import RAGPipeline
import utils.plot_generator as pg

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH, settings=Settings(allow_reset=True))

# Initialize session manager
session_manager = SessionManager()

# UI Configuration
st.set_page_config(
    page_title="Intelligent Data Analyzer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("📊 Intelligent Data Analyzer")
st.markdown("Upload your CSV or Excel file for advanced data analysis and visualization")

# --- Report Generation Function ---
def generate_report(df, file_name_prefix, report_type="profile"):
    """Generate reports in HTML, PDF, and DOCX formats."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_files = {}
    
    # Basic DataFrame info
    summary = f"Dataset: {file_name_prefix}\nRows: {df.shape[0]}\nColumns: {df.shape[1]}\n\n"
    columns_info = df.dtypes.to_string()
    stats = df.describe().to_string()
    
    # HTML Report
    html_path = os.path.join(PROFILES_DIR, f"{file_name_prefix}_profile_{timestamp}.html")
    html_content = f"""
    <html>
    <head><title>Data Profile Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #2c3e50; }}
        pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; }}
    </style>
    </head>
    <body>
        <h1>Data Profile Report: {file_name_prefix}</h1>
        <h2>Summary</h2>
        <pre>{summary}</pre>
        <h2>Column Types</h2>
        <pre>{columns_info}</pre>
        <h2>Basic Statistics</h2>
        <pre>{stats}</pre>
    </body>
    </html>
    """
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    report_files['html'] = html_path
    
    # DOCX Report
    docx_path = os.path.join(PROFILES_DIR, f"{file_name_prefix}_profile_{timestamp}.docx")
    doc = Document()
    doc.add_heading(f"Data Profile Report: {file_name_prefix}", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Column Types", level=1)
    doc.add_paragraph(columns_info)
    doc.add_heading("Basic Statistics", level=1)
    doc.add_paragraph(stats)
    doc.save(docx_path)
    report_files['docx'] = docx_path
    
    # PDF Report (LaTeX)
    latex_path = os.path.join(PROFILES_DIR, f"{file_name_prefix}_profile_{timestamp}.tex")
    latex_content = r"""
    \documentclass[a4paper,12pt]{article}
    \usepackage[utf8]{inputenc}
    \usepackage[T1]{fontenc}
    \usepackage{lmodern}
    \usepackage{geometry}
    \geometry{margin=1in}
    \usepackage{booktabs}
    \usepackage{longtable}
    \usepackage{pdflscape}
    \title{Data Profile Report: """ + file_name_prefix + r"""}
    \author{}
    \date{}
    \begin{document}
    \maketitle
    \section{Summary}
    \begin{verbatim}
    """ + summary.replace('$', '\$').replace('%', '\%') + r"""
    \end{verbatim}
    \section{Column Types}
    \begin{verbatim}
    """ + columns_info.replace('$', '\$').replace('%', '\%') + r"""
    \end{verbatim}
    \section{Basic Statistics}
    \begin{landscape}
    \begin{longtable}{lrrrrrrrr}
    """ + df.describe().to_latex().replace('$', '\$').replace('%', '\%') + r"""
    \end{longtable}
    \end{landscape}
    \end{document}
    """
    with open(latex_path, "w", encoding="utf-8") as f:
        f.write(latex_content)
    
    # Compile LaTeX to PDF
    try:
        import subprocess
        subprocess.run(["latexmk", "-pdf", "-quiet", latex_path], cwd=PROFILES_DIR, check=True)
        pdf_path = os.path.join(PROFILES_DIR, f"{file_name_prefix}_profile_{timestamp}.pdf")
        report_files['pdf'] = pdf_path
    except Exception as e:
        logging.error(f"Error compiling LaTeX to PDF: {e}")
        report_files['pdf'] = None
    
    return report_files

# --- Enhanced Dataset Management Functions ---
def find_dataset_by_name(search_term):
    """Find dataset by name using fuzzy matching"""
    datasets = session_manager.get_dataset_history()
    if not datasets:
        return None
    
    best_match = None
    best_score = 0
    
    for dataset in datasets:
        did, name, _, rows, cols, _ = dataset
        clean_name = os.path.splitext(name)[0].lower()
        search_clean = search_term.lower()
        
        if search_clean in clean_name or clean_name in search_clean:
            return dataset
        
        score = fuzz.partial_ratio(search_clean, clean_name)
        if score > best_score and score > 60:
            best_score = score
            best_match = dataset
    
    return best_match

def switch_to_dataset(dataset_info):
    """Switch to a specific dataset and update session"""
    try:
        did, file_name, _, rows, cols, _ = dataset_info
        file_path = os.path.join(DATA_DIR, file_name)
        
        if not os.path.exists(file_path):
            return False, f"Dataset file not found: {file_name}"
        
        data_loader = DataLoader()
        df = None
        
        if file_name.endswith('.csv'):
            df = data_loader.load_csv(file_path)
        elif file_name.endswith('.xlsx'):
            df = data_loader.load_excel(file_path)
        
        if df is not None:
            session_manager.set_dataframe(df, file_name)
            st.session_state.uploaded_file_name = file_name
            st.session_state.current_dataset_id = did
            
            st.session_state.rag_pipeline = RAGPipeline(
                df=df,
                chat_memory_manager=st.session_state.chat_memory_manager,
                temperature=st.session_state.llm_handler.temperature,
                top_p=st

.session_state.llm_handler.top_p
            )
            
            switch_message = f"✅ Successfully switched to dataset: {file_name} (ID: {did}) with {rows} rows and {cols} columns"
            st.session_state.chat_memory_manager.add_message("system", switch_message)
            
            return True, switch_message
        else:
            return False, f"Failed to load dataset: {file_name}"
    
    except Exception as e:
        logging.error(f"Error switching dataset: {e}")
        return False, f"Error switching dataset: {str(e)}"

def parse_dataset_switch_query(query):
    """Parse natural language queries for dataset switching"""
    query_lower = query.lower()
    
    switch_patterns = [
        r"switch to (.+?)(?:\s|$)",
        r"use (.+?)(?:\s|$)",
        r"load (.+?)(?:\s|$)",
        r"change to (.+?)(?:\s|$)",
        r"open (.+?)(?:\s|$)",
        r"analyze (.+?)(?:\s|$)",
        r"work with (.+?)(?:\s|$)",
        r"go to (.+?)(?:\s|$)"
    ]
    
    dataset_keywords = ["dataset", "data", "file", "table", "csv", "excel"]
    
    switch_intent = any(word in query_lower for word in ["switch", "use", "load", "change", "open", "analyze", "work with", "go to"])
    
    if switch_intent:
        for pattern in switch_patterns:
            match = re.search(pattern, query_lower)
            if match:
                search_term = match.group(1).strip()
                stop_words = ["the", "my", "dataset", "data", "file", "table", "previous", "old", "new", "last"]
                search_words = [word for word in search_term.split() if word not in stop_words]
                
                if search_words:
                    return " ".join(search_words)
    
    return None

def get_dataset_summary():
    """Get a summary of all available datasets"""
    datasets = session_manager.get_dataset_history()
    if not datasets:
        return "No datasets available. Upload a file first."
    
    summary = "**Available Datasets:**\n"
    current_id = st.session_state.get("current_dataset_id", None)
    
    for did, name, _, rows, cols, _ in datasets:
        status = " ← Current" if did == current_id else ""
        clean_name = os.path.splitext(name)[0]
        summary += f"- **{clean_name}** (ID: {did}): {rows} rows, {cols} columns{status}\n"
    
    return summary

def enhanced_query_processor(prompt):
    """Enhanced query processing with dataset switching capabilities"""
    dataset_search = parse_dataset_switch_query(prompt)
    
    if dataset_search:
        dataset_info = find_dataset_by_name(dataset_search)
        
        if dataset_info:
            success, message = switch_to_dataset(dataset_info)
            if success:
                clean_prompt = re.sub(
                    r"(switch to|use|load|change to|open|analyze|work with|go to)\s+[^.!?]*",
                    "",
                    prompt,
                    flags=re.IGNORECASE
                ).strip()
                
                if clean_prompt and len(clean_prompt) > 3:
                    try:
                        response = st.session_state.rag_pipeline.process_query(clean_prompt)
                        return response
                    except Exception as e:
                        return f"{message}\n\nError processing follow-up: {str(e)}"
                else:
                    return f"{message}\n\nWhat would you like to do with this dataset?"
            else:
                return f"Could not switch to dataset '{dataset_search}': {message}"
        else:
            return f"Dataset '{dataset_search}' not found.\n\n{get_dataset_summary()}"
    
    list_queries = ["list datasets", "show datasets", "what datasets", "available datasets", "my datasets"]
    if any(query in prompt.lower() for query in list_queries):
        return get_dataset_summary()
    
    if session_manager.is_data_loaded():
        try:
            response = st.session_state.rag_pipeline.process_query(prompt)
            return response
        except Exception as e:
            return f"Error processing your query: {str(e)}"
    else:
        return "Please upload a dataset first or switch to an existing dataset. Type 'list datasets' to see available options."

# --- Helper Functions ---
def get_file_download_link(file_path, file_label):
    """Generates a download link for a file."""
    try:
        if not os.path.exists(file_path):
            return f"<span style='color: red;'>File not found: {file_label}</span>"
        
        with open(file_path, "rb") as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        ext = os.path.splitext(file_path)[1][1:]
        return f'<a href="data:file/{ext};base64,{b64}" download="{os.path.basename(file_path)}">{file_label}</a>'
    except Exception as e:
        logging.error(f"Error creating download link: {e}")
        return f"<span style='color: red;'>Error creating download link</span>"

def display_plot_in_chat(plot_path):
    """Displays ONLY HTML plots directly in chat interface"""
    try:
        if not os.path.exists(plot_path):
            st.error(f"Plot file not found: {plot_path}")
            return False
        
        if plot_path.endswith('.html'):
            try:
                with open(plot_path, "r", encoding="utf-8") as f:
                    html_content = f.read()
                
                session_manager.store_plot(plot_path)
                components.html(html_content, height=600, scrolling=True)
                return True
            except Exception as e:
                st.error(f"Error rendering HTML plot: {str(e)}")
                st.markdown(get_file_download_link(plot_path, "Download Interactive Plot"), unsafe_allow_html=True)
                return False
        else:
            st.markdown(f"📊 **Plot generated:** `{os.path.basename(plot_path)}`")
            st.markdown(get_file_download_link(plot_path, f"Download {os.path.basename(plot_path)}"), unsafe_allow_html=True)
            return False
        
    except Exception as e:
        logging.error(f"Error displaying plot: {e}")
        st.error(f"Could not display plot: {str(e)}")
        return False

def get_recent_plots(seconds_ago=30):
    """Get ONLY HTML plots created within the last N seconds"""
    recent_plots = []
    current_time = datetime.now().timestamp()
    
    if os.path.exists(INTERACTIVE_PLOTS_DIR):
        for file in os.listdir(INTERACTIVE_PLOTS_DIR):
            if file.endswith('.html'):
                file_path = os.path.join(INTERACTIVE_PLOTS_DIR, file)
                if os.path.getmtime(file_path) > current_time - seconds_ago:
                    recent_plots.append(('HTML', file_path))
    
    return recent_plots

def display_chat_message(message):
    """Displays a single chat message with HTML report and plot handling"""
    with st.chat_message(message["role"]):
        content = message["content"]
        
        if message["role"] == "assistant":
            report_pattern = r"\[REPORT_(\w+)_PATH:([^\]]+)\]"
            plot_pattern = r"\[PLOT_(\w+)_PATH:([^\]]+)\]"
            
            reports = re.findall(report_pattern, content)
            plots = re.findall(plot_pattern, content)
            
            clean_content = re.sub(report_pattern, "", content)
            clean_content = re.sub(plot_pattern, "", content).strip()
            
            if clean_content:
                st.write(clean_content)
            
            if reports:
                st.subheader("Generated Reports:")
                for fmt, path in reports:
                    if os.path.exists(path):
                        if fmt.lower() == 'html':
                            try:
                                with open(path, "r", encoding="utf-8") as f:
                                    html_content = f.read()
                                components.html(html_content, height=800, scrolling=True)
                            except Exception as e:
                                st.error(f"Error rendering HTML report: {str(e)}")
                                st.markdown(get_file_download_link(path, f"Download HTML Report"), unsafe_allow_html=True)
                        else:
                            st.markdown(f"- **{fmt.upper()} Report**: `{os.path.basename(path)}`")
                            st.markdown(get_file_download_link(path, f"Download {fmt.upper()} Report"), unsafe_allow_html=True)
                    else:
                        st.warning(f"Report file not found: {path}")
            
            if "plots" in message and message["plots"]:
                st.subheader("📊 Generated Visualizations:")
                for plot_path in message["plots"]:
                    if plot_path.endswith('.html'):
                        display_plot_in_chat(plot_path)
        else:
            st.write(content)

def process_chat_with_html_detection(prompt, full_response):
    """Enhanced processing to auto-detect and display HTML plots and reports"""
    plot_keywords = ["plot", "chart", "graph", "histogram", "scatter", "bar", "pie", "heatmap", "visualization", "interactive", 'box', 'pairplot', 'line', 'violin', 'qq']
    
    if any(keyword in prompt.lower() for keyword in plot_keywords):
        time.sleep(2)
        recent_plots = get_recent_plots(60)
        if recent_plots and not re.search(r'\[PLOT_HTML_PATH:', full_response):
            st.subheader("📊 Auto-detected Interactive Visualizations:")
            for fmt, path in recent_plots:
                st.markdown(f"**{fmt} Interactive Plot:**")
                success = display_plot_in_chat(path)
                if success:
                    st.markdown(get_file_download_link(path, f"Download {fmt} Plot"), unsafe_allow_html=True)
    
    return full_response

# --- Sidebar ---
with st.sidebar:
    st.header("Upload Data & Settings")
    
    uploaded_file = st.file_uploader(
        "Upload a CSV or Excel file",
        type=SUPPORTED_FILE_TYPES,
        accept_multiple_files=False
    )
    
    if uploaded_file is not None:
        # Do not reset session on new upload to preserve chat history and UI state
        st.session_state.uploaded_file_name = uploaded_file.name
        
        file_path = os.path.join(DATA_DIR, uploaded_file.name)
        try:
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            data_loader = DataLoader()
            if uploaded_file.name.endswith('.csv'):
                df = data_loader.load_csv(file_path)
            elif uploaded_file.name.endswith('.xlsx'):
                df = data_loader.load_excel(file_path)

            if df is not None:
                session_manager.set_dataframe(df, uploaded_file.name)
                st.success(f"Loaded '{uploaded_file.name}'! Shape: {df.shape}")
                
                dataset_id = session_manager.record_dataset(uploaded_file.name, df)
                st.session_state.current_dataset_id = dataset_id
                
                st.session_state.rag_pipeline = RAGPipeline(
                    df=df,
                    chat_memory_manager=st.session_state.chat_memory_manager,
                    temperature=st.session_state.llm_handler.temperature,
                    top_p=st.session_state.llm_handler.top_p
                )
                
                init_msg = f"""
**I've loaded your data from {uploaded_file.name}** 🎉

Here's what you can do:

- **Dataset Management**: "list datasets", "switch to sales data", "use customer file"
- **Data Exploration**: "What columns do I have?", "Show me the first 10 rows"
- **Visualizations**: "Create a scatter plot of sales vs profit", "Show me a histogram of age"
- **Analysis**: "What are the key relationships?", "Give me analysis recommendations"
- **Reports**: "Generate a full data profile report"

**Dataset Switching Examples:**

- "Switch to my sales data and show top 10 products"
- "Use the customer dataset and create a pie chart of countries"
- "Load the previous file and analyze the correlations"

**Note:** Interactive HTML plots and reports are displayed directly in the chat. PDF and DOCX reports are available for download.

Try asking: "What insights can you give me about this data?"
"""
                # Append to existing messages instead of resetting
                if "messages" not in st.session_state:
                    st.session_state.messages = []
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": init_msg,
                    "plots": []
                })
                st.session_state.chat_memory_manager.add_message("assistant", init_msg)
            else:
                st.error("Failed to load DataFrame. Please check the file format.")
        except Exception as e:
            st.error(f"Error processing file: {e}")

    st.subheader("Dataset Management")
    datasets = session_manager.get_dataset_history()
    
    if datasets:
        current_id = st.session_state.get("current_dataset_id", None)
        
        st.write("**Available Datasets:**")
        for did, name, _, rows, cols, _ in datasets:
            status = " ← Current" if did == current_id else ""
            clean_name = os.path.splitext(name)[0]
            st.write(f"• **{clean_name}** (ID: {did}): {rows}×{cols}{status}")
        
        dataset_options = [f"{os.path.splitext(name)[0]} (ID: {did})" for did, name, _, _, _, _ in datasets]
        selected_dataset = st.selectbox("Quick Switch Dataset", dataset_options, key="quick_switch")
        
        if st.button("Switch Dataset", key="switch_btn"):
            dataset_id = int(selected_dataset.split('ID: ')[1].split(')')[0])
            dataset_info = next((d for d in datasets if d[0] == dataset_id), None)
            
            if dataset_info:
                success, message = switch_to_dataset(dataset_info)
                if success:
                    st.success(f"Switched to {dataset_info[1]}")
                    st.rerun()
                else:
                    st.error(message)
        
        st.info("💡 You can also switch datasets naturally in chat:\n'switch to sales data', 'use customer file', etc.")
    else:
        st.info("No datasets available. Upload a file to begin.")

    if session_manager.is_data_loaded():
        st.success("Data Loaded: Yes")
        df = session_manager.get_dataframe()
        current_file = st.session_state.get("uploaded_file_name", "Unknown")
        st.write(f"**Current File:** {current_file}")
        st.write(f"**Shape:** {df.shape[0]} rows, {df.shape[1]} columns")
        
        row_display = st.selectbox("Show rows:", [5, 10, 20, 30, 50], index=0)
        with st.expander(f"Data Preview (First {row_display} rows)"):
            st.dataframe(df.head(row_display))
    else:
        st.warning("No data loaded. Upload a file or switch to existing dataset.")

    st.markdown("---")
    st.header("Analysis Settings")
    
    if 'llm_handler' not in st.session_state:
        from utils.llm_handler import LLMHandler
        st.session_state.llm_handler = LLMHandler(temperature=0.7, top_p=0.95)
    
    current_temp = st.session_state.llm_handler.temperature
    current_top_p = st.session_state.llm_handler.top_p
    
    new_temp = st.slider("Temperature (Creativity)", 0.0, 1.0, current_temp, 0.01)
    new_top_p = st.slider("Top P (Diversity)", 0.0, 1.0, current_top_p, 0.01)
    
    if new_temp != current_temp or new_top_p != current_top_p:
        st.session_state.llm_handler.update_model_params(temperature=new_temp, top_p=new_top_p)
        if 'rag_pipeline' in st.session_state:
            st.session_state.rag_pipeline.llm_handler.update_model_params(temperature=new_temp, top_p=new_top_p)
        st.info("LLM parameters updated.")

    st.markdown("---")
    st.header("Actions")
    
    debug_mode = st.checkbox("🔧 Debug Mode", help="Show debug information for HTML plot generation")
    
    if st.button("🔄 Reset Session & Clear Data", help="Clears all loaded data, chat history, and generated reports"):
        session_manager.reset_session()
        st.success("Session reset complete! Upload a new file to continue.")
    
    if st.button("🔄 Refresh HTML Plots & Reports", help="Manually display the latest generated HTML content"):
        recent_plots = get_recent_plots(60)
        if recent_plots:
            st.success("HTML content refreshed!")
        else:
            st.info("No recent HTML content found.")

    st.markdown("---")
    st.header("Generated Reports & Visualizations")
    
    if st.button("🗑️ Clear All Plots & Reports", help="Remove all stored visualizations and reports"):
        if 'persistent_plots' in st.session_state:
            st.session_state.persistent_plots = []
            st.success("All content cleared!")
    
    st.subheader("Data Profile Reports")
    for fmt in REPORT_FORMATS:
        if os.path.exists(PROFILES_DIR):
            reports = [f for f in os.listdir(PROFILES_DIR) if f.endswith(f'.{fmt}')]
            if reports:
                for report in sorted(reports):
                    path = os.path.join(PROFILES_DIR, report)
                    st.markdown(f"**{fmt.upper()} Report:** {report}")
                    if fmt.lower() == 'html':
                        try:
                            with open(path, "r", encoding="utf-8") as f:
                                html_content = f.read()
                            components.html(html_content, height=800, scrolling=True)
                        except Exception as e:
                            st.error(f"Error rendering HTML report: {str(e)}")
                            st.markdown(get_file_download_link(path, f"Download {report}"), unsafe_allow_html=True)
                    else:
                        st.markdown(get_file_download_link(path, f"Download {report}"), unsafe_allow_html=True)
            else:
                st.info(f"No {fmt.upper()} profile reports generated yet.")
    
    st.subheader("Static Visualizations (PNG)")
    if os.path.exists(INSIGHTS_DIR):
        static_plots = [f for f in os.listdir(INSIGHTS_DIR) if f.endswith(('.png'))]
        if static_plots:
            for plot in sorted(static_plots):
                path = os.path.join(INSIGHTS_DIR, plot)
                st.image(path, caption=plot, use_column_width=True)
                st.markdown(get_file_download_link(path, f"Download {plot}"), unsafe_allow_html=True)
        else:
            st.info("No static visualizations generated yet.")
    
    st.subheader("Interactive Visualizations (HTML)")
    if os.path.exists(INTERACTIVE_PLOTS_DIR):
        interactive_plots = [f for f in os.listdir(INTERACTIVE_PLOTS_DIR) if f.endswith('.html')]
        if interactive_plots:
            for plot in sorted(interactive_plots):
                path = os.path.join(INTERACTIVE_PLOTS_DIR, plot)
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        html_content = f.read()
                    components.html(html_content, height=600, scrolling=True)
                except Exception as e:
                    st.error(f"Error rendering HTML plot: {str(e)}")
                    st.markdown(get_file_download_link(path, f"Download {plot}"), unsafe_allow_html=True)
        else:
            st.info("No interactive visualizations generated yet.")

# --- Chat Interface ---
st.header("Data Analysis Chat")
st.info("💡 Interactive HTML visualizations and reports will be displayed directly in the chat. PDF and DOCX reports are available for download.")

if "messages" not in st.session_state:
    st.session_state.messages = []
    
if "chat_memory_manager" not in st.session_state:
    from utils.memory_manager import ChatMemoryManager
    st.session_state.chat_memory_manager = ChatMemoryManager()

if 'persistent_plots' not in st.session_state:
    st.session_state.persistent_plots = []

for message in st.session_state.messages:
    display_chat_message(message)

if prompt := st.chat_input("Ask me about your data or switch datasets..."):
    st.session_state.messages.append({
        "role": "user",
        "content": prompt
    })
    st.session_state.chat_memory_manager.add_message("user", prompt)
    
    with st.chat_message("user"):
        st.write(prompt)
    
    prev_plot_count = len(st.session_state.persistent_plots)
    
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        full_response = ""
        
        if prompt.strip().lower() == "/reset":
            session_manager.reset_session()
            full_response = "Session has been reset. Please upload a new file."
        elif prompt.strip().lower() == "/profile":
            if session_manager.is_data_loaded():
                df = session_manager.get_dataframe()
                prefix = os.path.splitext(st.session_state.uploaded_file_name)[0]
                report_files = generate_report(df, prefix)
                full_response = "Generated profile reports:\n"
                for fmt, path in report_files.items():
                    if path and os.path.exists(path):
                        full_response += f"[REPORT_{fmt.upper()}_PATH:{path}] "
                    else:
                        full_response += f"Failed to generate {fmt.upper()} report. "
            else:
                full_response = "Please upload data first to generate a profile."
        elif prompt.strip().lower() in ["/insights", "insights"]:
            if session_manager.is_data_loaded():
                tool_output = st.session_state.rag_pipeline.tools[4]()
                full_response = tool_output
            else:
                full_response = "Please upload data first to get insights."
        elif prompt.strip().lower() in ["/help", "help"]:
            full_response = """
            **Available Commands:**

            - `/reset`: Clear session and reset application
            - `/profile`: Generate comprehensive data profile report (HTML, PDF, DOCX)
            - `/insights`: Get key insights and recommendations
            - `help`: Show this help message
            
            **Natural Dataset Switching:**

            - "switch to sales data" or "use customer file"
            - "load my previous dataset" or "change to inventory file"
            - "work with the user data" or "analyze the product table"
            - "list datasets" or "show available datasets"
            
            **Combined Actions:**

            - "switch to sales data and show top 10 products"
            - "use customer file and create a pie chart of countries"
            - "load inventory data and analyze stock levels"
            
            **Analysis Options:**

            - **Basic Info**: "What columns do I have?", "Show me the first 20 rows"
            - **Interactive Visualizations**: "Create an interactive scatter plot", "Show histogram"
            - **Data Insights**: "What are the key patterns?", "Any data quality issues?"
            - **Machine Learning**: "What models could I build?", "Predict customer behavior"
            
            **Visualization Examples:**

            - "Create interactive scatter plot of sales vs profit"
            - "Show me interactive histogram of age distribution"
            - "Generate correlation heatmap for numerical columns"
            - "Make an interactive bar chart of top products"
            
            **Dataset Management:**

            - Datasets are automatically saved when uploaded
            - Switch between datasets using natural language
            - All analysis works with the currently active dataset
            - Use fuzzy matching: "sales" matches "sales_data_2024.csv"
            """
        else:
            full_response = enhanced_query_processor(prompt)
        
        message_placeholder.markdown(full_response)
        
        if debug_mode:
            st.write("🔧 **Debug Info:**")
            st.write(f"Current dataset: {st.session_state.get('uploaded_file_name', 'None')}")
            st.write(f"Dataset ID: {st.session_state.get('current_dataset_id', 'None')}")
            st.write(f"Data loaded: {session_manager.is_data_loaded()}")
            st.write(f"Previous plot count: {prev_plot_count}")
            st.write(f"Current plot count: {len(st.session_state.persistent_plots)}")
        
        full_response = process_chat_with_html_detection(prompt, full_response)
        
        new_plots = st.session_state.persistent_plots[prev_plot_count:]
        
        st.session_state.messages.append({
            "role": "assistant",
            "content": full_response,
            "plots": new_plots
        })
        st.session_state.chat_memory_manager.add_message("assistant", full_response)